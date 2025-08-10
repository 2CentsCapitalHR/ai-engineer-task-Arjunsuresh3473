import os
import io
import json
import streamlit as st
from docx import Document as DocxDocument
from docx.shared import RGBColor
from unidecode import unidecode
from typing import List, Dict, Any, Tuple, Optional

# RAG & LLM dependencies
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDoc
from langchain_community.vectorstores import FAISS
from langchain.embeddings.base import Embeddings
import google.generativeai as genai
from google.generativeai.types import GenerationConfig

# ---------------------------
# CONFIG
# ---------------------------
REFERENCE_DIR = os.path.join("data", "reference")
FAISS_INDEX_DIR = os.path.join("data", "faiss_index")

MODEL_NAME = "models/embedding-001"
GENERATION_MODEL = "gemini-1.5-flash"
MAX_CHARS = 2500

INCORPORATION_CHECKLIST = [
    "Articles of Association (AoA)",
    "Memorandum of Association (MoA/MoU)",
    "Incorporation Application Form",
    "UBO Declaration Form",
    "Register of Members and Directors",
]

DOC_TYPE_KEYWORDS = {
    "articles of association": "Articles of Association (AoA)",
    "aoa": "Articles of Association (AoA)",
    "memorandum of association": "Memorandum of Association (MoA/MoU)",
    "moa": "Memorandum of Association (MoA/MoU)",
    "mou": "Memorandum of Association (MoA/MoU)",
    "incorporation application": "Incorporation Application Form",
    "ubo declaration": "UBO Declaration Form",
    "register of members and directors": "Register of Members and Directors",
}

# ---------------------------
# Utilities
# ---------------------------
def normalize_text(s: str) -> str:
    return unidecode((s or "")).lower()

def safe_first(li: List[Any], default=None):
    return li[0] if li else default

# Instead of docx comments, add a red flag marker text in red color
def mark_paragraph_as_redflag(paragraph, text):
    run = paragraph.add_run(f"\n⚠️ RED FLAG: {text}")
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Bright Red color

# ---------------------------
# Gemini embeddings adapter inheriting from LangChain Embeddings
# ---------------------------
class GeminiEmbeddingsForLangChain(Embeddings):
    def __init__(self, api_key: str):
        genai.configure(api_key=api_key)
        self.model_name = MODEL_NAME
        self.max_chars = MAX_CHARS

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        return [self._embed_content(text) for text in texts]

    def embed_query(self, text: str) -> List[float]:
        return self._embed_content(text)

    def _embed_content(self, text: str) -> List[float]:
        text = text[:self.max_chars]
        try:
            response = genai.embed_content(model=self.model_name, content=text)
            return response["embedding"]
        except Exception as e:
            st.error(f"Embedding failed: {str(e)}")
            return []

# ---------------------------
# Gemini generation helper
# ---------------------------
def gemini_generate(prompt: str, api_key: str, max_tokens: int = 300) -> str:
    try:
        genai.configure(api_key=api_key)
        model = genai.GenerativeModel(GENERATION_MODEL)
        response = model.generate_content(
            prompt[:MAX_CHARS],
            generation_config=GenerationConfig(
                max_output_tokens=max_tokens
            )
        )
        # Use response.text if available; fallback
        if response.candidates and response.text:
            return response.text
        return "No response generated"
    except Exception as e:
        st.error(f"Generation failed: {str(e)}")
        return "Error generating response"

# ---------------------------
# FAISS vector store management
# ---------------------------
@st.cache_resource(show_spinner=False)
def build_or_load_faiss(api_key: str, reference_dir: str = REFERENCE_DIR, index_path: str = FAISS_INDEX_DIR):
    embeddings = GeminiEmbeddingsForLangChain(api_key)

    if os.path.exists(index_path) and os.path.isdir(index_path):
        try:
            return FAISS.load_local(index_path, embeddings, allow_dangerous_deserialization=True)
        except Exception as e:
            st.warning(f"Couldn't load existing index, rebuilding: {str(e)}")
    
    docs: List[LangchainDoc] = []
    if not os.path.exists(reference_dir):
        st.error("Reference data directory 'data/reference' not found.")
        return FAISS.from_documents([LangchainDoc(page_content="ADGM reference laws.", metadata={})], embeddings)

    for fname in sorted(os.listdir(reference_dir)):
        fpath = os.path.join(reference_dir, fname)
        if not os.path.isfile(fpath):
            continue
            
        text = ""
        try:
            if fname.lower().endswith(".docx"):
                doc = DocxDocument(fpath)
                text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            elif fname.lower().endswith(".txt"):
                with open(fpath, "r", encoding="utf-8") as f:
                    text = f.read()
            else:
                continue
            
            if text.strip():
                docs.append(LangchainDoc(page_content=text, metadata={"source": fname}))
        except Exception as e:
            st.warning(f"Error processing {fname}: {str(e)}")

    if not docs:
        st.warning("No reference documents found. RAG will not function.")
        return FAISS.from_documents([LangchainDoc(page_content="ADGM reference laws.", metadata={})], embeddings)

    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
    split_docs = splitter.split_documents(docs)
    
    vector_store = FAISS.from_documents(split_docs, embeddings)
    
    try:
        os.makedirs(index_path, exist_ok=True)
        vector_store.save_local(index_path)
    except Exception as e:
        st.warning(f"Couldn't save index: {str(e)}")
    
    return vector_store

# ---------------------------
# Core Logic
# ---------------------------
def identify_document_type(doc_text: str) -> Optional[str]:
    for keyword, doc_type in DOC_TYPE_KEYWORDS.items():
        if keyword in normalize_text(doc_text):
            return doc_type
    return None

def analyze_document(doc_text: str, doc_type: str, vector_store) -> Tuple[List[Dict[str, Any]], str]:
    issues_found = []
    
    # RAG-based Red Flag Detection queries
    red_flag_queries = [
        f"Does this {doc_type} text contain any clauses that are not compliant with ADGM regulations?",
        f"Are there any references to jurisdictions other than ADGM within this {doc_type}?",
        f"Is the language in this {doc_type} clear and legally binding, or is it ambiguous?",
    ]

    reviewed_doc_text = doc_text
    
    for query in red_flag_queries:
        context_docs = vector_store.similarity_search(query, k=3)
        context = " ".join([doc.page_content for doc in context_docs])
        
        prompt = f"""
        You are an ADGM-compliant corporate legal agent.
        Your task is to review the following legal document text:
        ---
        {doc_text[:MAX_CHARS]}
        ---
        
        Using the following ADGM laws and regulations as context:
        ---
        {context[:MAX_CHARS]}
        ---
        
        Please identify any legal red flags or inconsistencies based on the query: '{query}'.
        Focus on issues like incorrect jurisdiction, ambiguous language, or non-compliance with ADGM law.
        If a red flag is found, describe the issue, suggest a correction, and, if possible, cite the relevant ADGM rule.
        Return the response as a simple paragraph. If no issues are found, just say so.
        """
        
        response = gemini_generate(prompt, st.session_state.api_key, max_tokens=500)
        
        if "no issues" not in response.lower() and "no red flags" not in response.lower():
            # Very simple parsing to get first sentence as issue, rest as suggestion
            sentences = response.strip().split('.')
            issue = sentences[0].strip() if sentences else response.strip()
            suggestion = sentences[1].strip() if len(sentences) > 1 else ""
            
            issue_summary = {
                "document": doc_type,
                "section": "General",  # Placeholder
                "issue": issue,
                "severity": "High",
                "suggestion": suggestion
            }
            issues_found.append(issue_summary)
    
    return issues_found, reviewed_doc_text

def process_uploaded_files(uploaded_files: List[io.BytesIO], api_key: str):
    if not uploaded_files:
        st.warning("Please upload at least one document.")
        return

    st.session_state.uploaded_file_names = [f.name for f in uploaded_files]
    uploaded_docs = {}
    
    for uploaded_file in uploaded_files:
        try:
            doc = DocxDocument(uploaded_file)
            doc_text = "\n".join(p.text for p in doc.paragraphs)
            doc_type = identify_document_type(doc_text)
            
            uploaded_docs[uploaded_file.name] = {
                "type": doc_type,
                "text": doc_text,
                "document_object": doc,
            }
        except Exception as e:
            st.error(f"Failed to process {uploaded_file.name}: {e}")
            return
            
    st.info("Step 1: Identifying legal process and checking document checklist...")
    process = "Company Incorporation" 
    uploaded_doc_types = [doc_data['type'] for doc_data in uploaded_docs.values() if doc_data['type']]
    
    missing_docs = [doc for doc in INCORPORATION_CHECKLIST if doc not in uploaded_doc_types]
    
    if missing_docs:
        missing_doc_str = ", ".join(missing_docs)
        st.warning(f"Missing documents: {missing_doc_str}")
    
    st.info("Step 2: Analyzing each document for red flags and inconsistencies...")
    
    vector_store = build_or_load_faiss(api_key)
    all_issues = []
    
    for filename, doc_data in uploaded_docs.items():
        if not doc_data['type']:
            st.warning(f"Could not identify the document type for '{filename}'. Skipping analysis.")
            continue
        
        with st.spinner(f"Analyzing {filename} for red flags..."):
            issues, reviewed_text = analyze_document(doc_data['text'], doc_data['type'], vector_store)
            all_issues.extend(issues)
            
            # Add red flag markers in the docx (append red text at the end of first paragraph)
            for issue in issues:
                if doc_data['document_object'].paragraphs:
                    mark_paragraph_as_redflag(doc_data['document_object'].paragraphs[0], issue['issue'])
            
            # Save the reviewed doc
            output_path = f"reviewed_{filename}"
            doc_data['document_object'].save(output_path)
            st.success(f"Review complete for {filename}. Download the reviewed file below.")
            with open(output_path, "rb") as f:
                st.download_button(
                    label=f"Download Reviewed {filename}",
                    data=f,
                    file_name=output_path,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
    st.info("Step 3: Generating final structured report...")
    structured_report = {
        "process": process,
        "documents_uploaded": len(uploaded_doc_types),
        "required_documents": len(INCORPORATION_CHECKLIST),
        "missing_documents": missing_docs,
        "issues_found": all_issues
    }
    
    st.json(structured_report)
    
    json_output = json.dumps(structured_report, indent=4)
    st.download_button(
        label="Download JSON Report",
        data=json_output,
        file_name="ADGM_Review_Report.json",
        mime="application/json"
    )

def main():
    st.set_page_config(page_title="ADGM Document Review", layout="wide")
    st.title("ADGM Corporate Document Review System")
    
    st.markdown("Please enter your Gemini API key to begin.")
    api_key = st.text_input("Gemini API Key", type="password")
    if api_key:
        st.session_state.api_key = api_key
        genai.configure(api_key=api_key)
        
        try:
            with st.spinner("Initializing AI components..."):
                _ = build_or_load_faiss(api_key)
            st.success("System ready! Now upload your documents.")
        except Exception as e:
            st.error(f"Initialization failed: {e}")
            st.stop()
            
        st.markdown("---")
        
        uploaded_files = st.file_uploader(
            "Upload your ADGM-related .docx documents",
            type="docx",
            accept_multiple_files=True
        )
        
        if uploaded_files:
            if st.button("Review Documents"):
                process_uploaded_files(uploaded_files, api_key)
    else:
        st.warning("Please enter your Gemini API key to proceed.")

if __name__ == "__main__":
    main()
